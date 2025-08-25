"""
File parsing service for extracting text from PDF and DOC files
"""
import os
import PyPDF2
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileParser:
    """Service for parsing resume files and extracting text content"""
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """
        Extract text from PDF file using PyPDF2
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            tuple: (success: bool, text: str, error: str)
        """
        try:
            text = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    return False, "", "PDF is password protected and cannot be processed"
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        text += page_text + "\n"
                
                # Clean up the text
                text = FileParser._clean_extracted_text(text)
                
                if not text.strip():
                    return False, "", "No readable text found in PDF. The file might be image-based or corrupted."
                
                logger.info(f"Successfully extracted {len(text)} characters from PDF: {file_path}")
                return True, text, ""
                
        except Exception as e:
            error_msg = f"Error parsing PDF file: {str(e)}"
            logger.error(error_msg)
            return False, "", error_msg
    
    @staticmethod
    def extract_text_from_docx(file_path_or_file):
        """
        Extract text from DOCX file using python-docx

        Args:
            file_path_or_file: Either a file path (str) or a file object

        Returns:
            tuple: (success: bool, text: str, error: str)
        """
        try:
            # Handle both file paths and file objects
            if hasattr(file_path_or_file, 'read'):
                # It's a file object (from Flask request)
                doc = Document(file_path_or_file)
                source_info = "uploaded file"
            else:
                # It's a file path
                doc = Document(file_path_or_file)
                source_info = file_path_or_file

            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Clean up the text
            text = FileParser._clean_extracted_text(text)
            
            if not text.strip():
                return False, "", "No readable text found in DOCX file."
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX: {source_info}")
            return True, text, ""
            
        except Exception as e:
            error_msg = f"Error parsing DOCX file: {str(e)}"
            logger.error(error_msg)
            return False, "", error_msg
    
    @staticmethod
    def extract_text_from_doc(file_path):
        """
        Extract text from DOC file (legacy format)
        Note: This is a simplified implementation. For production, consider using python-docx2txt or antiword
        
        Args:
            file_path (str): Path to the DOC file
            
        Returns:
            tuple: (success: bool, text: str, error: str)
        """
        try:
            # For now, we'll return an error for DOC files and suggest conversion
            # In production, you might want to use python-docx2txt or antiword
            return False, "", "Legacy DOC format is not supported. Please convert to DOCX or PDF format."
            
        except Exception as e:
            error_msg = f"Error parsing DOC file: {str(e)}"
            logger.error(error_msg)
            return False, "", error_msg
    
    @staticmethod
    def parse_resume_file(file_path, file_type):
        """
        Parse resume file based on its type
        
        Args:
            file_path (str): Path to the file
            file_type (str): Type of file (pdf, doc, docx)
            
        Returns:
            tuple: (success: bool, text: str, error: str)
        """
        if not os.path.exists(file_path):
            return False, "", "File not found"
        
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return FileParser.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return FileParser.extract_text_from_docx(file_path)
        elif file_type == 'doc':
            return FileParser.extract_text_from_doc(file_path)
        else:
            return False, "", f"Unsupported file type: {file_type}"
    
    @staticmethod
    def _clean_extracted_text(text):
        """
        Clean and normalize extracted text, removing problematic characters

        Args:
            text (str): Raw extracted text

        Returns:
            str: Cleaned text safe for database storage
        """
        if not text:
            return ""

        import re

        # Step 1: Remove null bytes and other problematic characters
        # Remove null bytes (0x00) that cause PostgreSQL errors
        text = text.replace('\x00', '')

        # Remove other control characters except newlines, tabs, and carriage returns
        text = re.sub(r'[\x01-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

        # Step 2: Normalize unicode characters
        import unicodedata
        text = unicodedata.normalize('NFKD', text)

        # Step 3: Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line:  # Only keep non-empty lines
                # Remove excessive spaces within the line
                line = re.sub(r' +', ' ', line)
                cleaned_lines.append(line)

        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)

        # Step 4: Ensure the text is not too long (database field limits)
        max_length = 50000  # Reasonable limit for resume text
        if len(cleaned_text) > max_length:
            cleaned_text = cleaned_text[:max_length] + "... [truncated]"

        return cleaned_text
    
    @staticmethod
    def get_file_info(file_path):
        """
        Get basic information about a file
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: File information
        """
        try:
            if not os.path.exists(file_path):
                return None
            
            stat = os.stat(file_path)
            
            return {
                'size': stat.st_size,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'exists': True
            }
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return None
    
    @staticmethod
    def validate_file_type(filename):
        """
        Validate if file type is supported
        
        Args:
            filename (str): Name of the file
            
        Returns:
            tuple: (is_valid: bool, file_type: str, error: str)
        """
        if not filename:
            return False, "", "No filename provided"
        
        # Get file extension
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ""
        
        supported_types = ['pdf', 'doc', 'docx']
        
        if file_ext not in supported_types:
            return False, "", f"Unsupported file type. Supported types: {', '.join(supported_types)}"
        
        return True, file_ext, ""
